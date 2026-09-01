# -*- coding: utf-8 -*-
"""Repair a paragraph whose trailing text was lost behind a REF field.

A cross-reference paragraph is [tab][REF field -> "Figure 40"][the sentence].
Writing replacement text into the paragraph's first w:t lands it inside the
field result, which Word then discards the next time it updates fields. The
sentence after the field end is the part that must be edited; this restores it
from the backup, with the migration clause already removed.
"""
from __future__ import annotations

import copy
import sys

import docx
from docx.oxml.ns import qn

BACKUP = "built-backup-before-migrations-removal.docx"
ANCHOR = "presents the entity relationship diagram of the deployed database"

OLD = ("a table added by a later migration appears in the diagram automatically rather than "
       "being forgotten.")
NEW = ("every table, column, key and relationship shown in it is read from the database itself.")


def tail_runs(p):
    """The runs that follow the last fldChar end - the paragraph's own text."""
    runs = p._p.findall(qn("w:r"))
    last_end = -1
    for i, r in enumerate(runs):
        fld = r.find(qn("w:fldChar"))
        if fld is not None and fld.get(qn("w:fldCharType")) == "end":
            last_end = i
    return runs[last_end + 1:]


def main() -> int:
    path = sys.argv[1]

    src = docx.Document(BACKUP)
    good = next(p for p in src.paragraphs if ANCHOR in p.text)
    good_tail = tail_runs(good)
    if not good_tail:
        raise SystemExit("backup paragraph has no text after the field")
    original = "".join(t.text or "" for r in good_tail for t in r.findall(qn("w:t")))
    if OLD not in original:
        raise SystemExit(f"clause not found in backup tail: {original[:120]!r}")
    fixed = original.replace(OLD, NEW)

    doc = docx.Document(path)
    target = next(p for p in doc.paragraphs
                  if p.text.strip() == "Figure 40" and p._p.find(qn("w:fldChar")) is None
                  or (p.text.strip() == "Figure 40"))
    # drop whatever trails the field now, then re-attach a single clean run
    for r in tail_runs(target):
        target._p.remove(r)
    run = copy.deepcopy(good_tail[0])
    for t in run.findall(qn("w:t")):
        run.remove(t)
    t = run.makeelement(qn("w:t"), {})
    t.text = fixed
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    target._p.append(run)

    doc.save(path)
    chk = docx.Document(path)
    out = next(p for p in chk.paragraphs if ANCHOR in p.text)
    print("repaired paragraph now reads:\n")
    print("  " + out.text.strip()[:300] + " ...")
    print(f"\n  length: {len(out.text.split())} words")
    print(f"  migration mentioned: {'migrat' in out.text.lower()}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
