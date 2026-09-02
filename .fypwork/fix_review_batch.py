# -*- coding: utf-8 -*-
"""The reviewer's findings, fixed in place.

Covers the numbering gap in Chapter 3, the self-referencing table of contents,
the references heading and its missing page break, the two chapter openings held
in position by blank paragraphs, the cover title, three alphabetical breaks and
one corrupted entry in the reference list, three wrong or duplicated captions,
the abstract's plan-era wording, and a heading that ends in a colon.
"""
from __future__ import annotations

import re
import sys

import docx
from docx.oxml.ns import qn

# ── Chapter 3: close the gap left where 3.3 should be ────────────────────
# Applied in ascending order so a new number never collides with a live one.
RENUMBER = [
    ("3.4 Methodology Discussion", "3.3 Methodology Discussion"),
    ("3.5 Selected methodology", "3.4 Selected methodology"),
    ("3.6 Activities and Processes", "3.5 Activities and Processes"),
    ("3.7 Data Gathering Design", "3.6 Data Gathering Design"),
    ("3.7.1 Data Collection Method", "3.6.1 Data Collection Method"),
    ("3.7.2 Target Participants", "3.6.2 Target Participants"),
    ("3.7.3 Sample Size and Distribution", "3.6.3 Sample Size and Distribution"),
    ("3.7.4 Survey Structure and Design", "3.6.4 Survey Structure and Design"),
    ("3.7.5 Considerations", "3.6.5 Considerations"),
    ("3.7.6 Purpose of Data Collection", "3.6.6 Purpose of Data Collection"),
    ("3.8 Analysis", "3.7 Analysis"),
    ("3.8.1 Quantitative Data Presentation", "3.7.1 Quantitative Data Presentation"),
    ("3.8.2 Summary of Findings", "3.7.2 Summary of Findings"),
    ("3.9 Summary", "3.8 Summary"),
]

# 3.8.1 and 3.8.2 were Heading 2, level with their own parent. They are children.
DEMOTE = ["3.7.1 Quantitative Data Presentation", "3.7.2 Summary of Findings"]

# body cross-references that name a renumbered section
CROSSREF = [("Section 3.7.4", "Section 3.6.4"), ("Section 3.5", "Section 3.4")]

TITLE_OLD = ("Developing a Centralized Role Based Event Coordination Platform for "
             "Multi Departments University Workflow Management")
TITLE_NEW = ("Developing A Centralized Role Based Event Coordination Platform For "
             "Multi Departments University Workflow Management")

ABSTRACT_OLD = ("The preliminary findings indicate that current systems fail to support "
                "independent departmental workflows within a unified platform, creating "
                "significant operational gaps.")
ABSTRACT_NEW = ("The findings indicate that existing systems fail to support independent "
                "departmental workflows within a unified platform, creating significant "
                "operational gaps.")

CAPTIONS = [
    ("Figure 132", "Proposal Form (PPF).", "Ethics Forms — Participant Confidentiality."),
    ("Figure 136", "Ethics Forms — Target Participants.", "Ethics Forms — Support Information."),
]

# reference list: (entry to move, entry it must precede)
REORDER = [("Ahmad, T., & Van Looy, A. (2020)", "Alami, A., & Krancher, O. (2022)"),
           ("Anderer, S., Kempter, T.", "Arriesgado, J., Calaguian, J."),
           ("ISGlobal. (n.d.)", "Ismail, R., Safieddine, F.")]

STRAY_DOI = "https://doi.org/10.1007/978-3-642-12636-9_9\t"

GITHUB_SENTENCE = ("The complete source code for the delivered platform is published at "
                   "https://github.com/osama583/EMS.")


# ── xml helpers ──────────────────────────────────────────────────────────
def text_nodes(p):
    """w:t nodes outside any field result."""
    out, depth = [], 0
    for r in p._p.findall(qn("w:r")):
        for node in r:
            if node.tag == qn("w:fldChar"):
                kind = node.get(qn("w:fldCharType"))
                depth += 1 if kind == "begin" else (-1 if kind == "end" else 0)
                depth = max(0, depth)
            elif node.tag == qn("w:t") and depth == 0:
                out.append(node)
    return out


def replace_in(p, old: str, new: str) -> bool:
    nodes = text_nodes(p)
    full = "".join(n.text or "" for n in nodes)
    at = full.find(old)
    if at < 0:
        return False
    end, pos, payload = at + len(old), 0, new
    for n in nodes:
        t = n.text or ""
        a, b = pos, pos + len(t)
        pos = b
        if b <= at or a >= end:
            continue
        lo, hi = max(at, a) - a, min(end, b) - a
        n.text = t[:lo] + payload + t[hi:]
        n.set(qn("xml:space"), "preserve")
        payload = ""
    return True


def get_ppr(p):
    ppr = p._p.find(qn("w:pPr"))
    if ppr is None:
        ppr = p._p.makeelement(qn("w:pPr"), {})
        p._p.insert(0, ppr)
    return ppr


def page_break_before(p) -> None:
    ppr = get_ppr(p)
    if ppr.find(qn("w:pageBreakBefore")) is not None:
        return
    node = ppr.makeelement(qn("w:pageBreakBefore"), {})
    style = ppr.find(qn("w:pStyle"))
    ppr.insert(list(ppr).index(style) + 1 if style is not None else 0, node)


def drop_from_toc(p) -> None:
    """Keep the look, lose the outline level, so the field stops indexing it."""
    ppr = get_ppr(p)
    if ppr.find(qn("w:outlineLvl")) is not None:
        ppr.find(qn("w:outlineLvl")).set(qn("w:val"), "9")
        return
    node = ppr.makeelement(qn("w:outlineLvl"), {qn("w:val"): "9"})
    rpr = ppr.find(qn("w:rPr"))
    ppr.insert(list(ppr).index(rpr) if rpr is not None else len(ppr), node)


CARRIES = (qn("a:blip"), qn("w:fldChar"), qn("w:instrText"),
           qn("w:bookmarkStart"), qn("w:sdt"))


def is_spacer(p) -> bool:
    """Empty of text AND of anything structural.

    A paragraph can look blank and still close a field: the List of Figures and
    List of Tables both end in one, and deleting it destroys the whole list.
    """
    if p.text.strip():
        return False
    return not any(p._p.findall(".//" + tag) or p._p.findall(tag) for tag in CARRIES)


def blanks_before(paras, idx) -> list:
    out = []
    j = idx - 1
    while j >= 0 and is_spacer(paras[j]):
        out.append(paras[j])
        j -= 1
    return out


def main() -> int:
    path = sys.argv[1]
    doc = docx.Document(path)
    done = []

    def para(pred):
        hits = [p for p in doc.paragraphs if pred(p)]
        if len(hits) != 1:
            raise SystemExit(f"expected 1 paragraph, got {len(hits)}")
        return hits[0]

    # 3 — Chapter 3 numbering
    for old, new in RENUMBER:
        p = para(lambda q, o=old: q.style.name.startswith("Heading") and q.text.strip().startswith(o))
        replace_in(p, old, new)
    h3 = doc.styles["Heading 3"]
    for target in DEMOTE:
        para(lambda q, t=target: q.text.strip().startswith(t)).style = h3
    for old, new in CROSSREF:
        for p in doc.paragraphs:
            if old in p.text and not p.style.name.startswith("Heading"):
                replace_in(p, old, new)
    done.append(f"Chapter 3 renumbered 3.4-3.9 -> 3.3-3.8, {len(DEMOTE)} sub-headings demoted")

    # 6 — table of contents indexing itself
    toc_heads = [p for p in doc.element.body.iter(qn("w:p"))
                 if "".join(t.text or "" for t in p.iter(qn("w:t"))).strip() == "TABLE OF CONTENTS"]
    for element in toc_heads:
        drop_from_toc(type("P", (), {"_p": element})())
    done.append(f"TABLE OF CONTENTS removed from its own field ({len(toc_heads)})")

    # 5 + 7 — references heading and its page break
    ref = para(lambda q: q.text.strip().startswith("REFERENCE"))
    replace_in(ref, "REFERENCE:", "REFERENCES")
    page_break_before(ref)
    done.append("REFERENCE: -> REFERENCES, on a new page")

    # 8 — blank paragraphs replaced by real page breaks
    for name in ("LIST OF TABLES", "CHAPTER 1: INTRODUCTION"):
        paras = doc.paragraphs
        idx = next(i for i, p in enumerate(paras) if p.text.strip() == name)
        gap = blanks_before(paras, idx)
        for blank in gap:
            blank._p.getparent().remove(blank._p)
        page_break_before(paras[idx])
        done.append(f"{name}: {len(gap)} blank paragraphs -> page break")

    # 9 + 10 — cover title
    cover = para(lambda q: q.text.strip().endswith(TITLE_OLD + "."))
    replace_in(cover, TITLE_OLD + ".", TITLE_NEW)
    for p in doc.paragraphs:
        if TITLE_OLD in p.text:
            replace_in(p, TITLE_OLD, TITLE_NEW)
    done.append("title: trailing stop removed, 'a'/'for' capitalised on cover and declaration")

    # 11 — corrupted reference entry
    broken = para(lambda q: STRAY_DOI in q.text)
    replace_in(broken, STRAY_DOI, "")
    done.append("stray DOI removed from the Mualla entry")

    # 12 — alphabetical order
    for mover, before in REORDER:
        a = para(lambda q, m=mover: q.text.strip().startswith(m))
        b = para(lambda q, t=before: q.text.strip().startswith(t))
        b._p.addprevious(a._p)
    done.append(f"{len(REORDER)} reference entries reordered")

    # 14 + 15 — captions
    for label, old, new in CAPTIONS:
        p = para(lambda q, l=label, o=old: q.style.name == "Caption"
                 and q.text.strip().startswith(l) and o in q.text)
        replace_in(p, old, new)
    done.append(f"{len(CAPTIONS)} captions corrected")

    # 19 — stray space before a caption colon
    fixed = 0
    for p in doc.paragraphs:
        if p.style.name != "Caption" or "Similarity Report" not in p.text:
            continue
        for n in text_nodes(p):
            if n.text and n.text.lstrip().startswith(":") and n.text != n.text.lstrip():
                n.text = n.text.lstrip()
                n.set(qn("xml:space"), "preserve")
                fixed += 1
            elif n.text and n.text.strip() == "" :
                n.text = ""
                fixed += 1
    done.append(f"caption colon spacing normalised ({fixed} run(s))")

    # 25 — abstract
    replace_in(para(lambda q: ABSTRACT_OLD in q.text), ABSTRACT_OLD, ABSTRACT_NEW)
    done.append("abstract no longer calls the findings preliminary")

    # 26 — heading colon
    replace_in(para(lambda q: q.text.strip() == "1.2.1 The Core Problem:"),
               "1.2.1 The Core Problem:", "1.2.1 The Core Problem")
    done.append("1.2.1 heading colon removed")

    # Appendix F — the link sentence duplicates the heading
    link = para(lambda q: GITHUB_SENTENCE in q.text)
    link._p.getparent().remove(link._p)
    done.append("Appendix F link sentence deleted (heading already carries it)")

    doc.save(path)
    for line in done:
        print(f"  {line}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
