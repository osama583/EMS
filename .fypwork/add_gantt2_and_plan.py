# -*- coding: utf-8 -*-
"""Task 2: add the FYP Part 2 Gantt chart and extend the project plan.

Nothing existing is replaced. Figures 140 and 141 stay exactly where they are,
including the two IR-era labels the supervisor accepted as correct for the time
they were written, and every one of the 43 project plan rows is left untouched.
This appends a second pair of Gantt figures after them and a block of new plan
rows covering the work that produced Chapters 4 to 6.

Captions carry a live SEQ field, so the new figures number themselves and reach
the List of Figures without any number being typed by hand.
"""
from __future__ import annotations

import copy
import os
import sys

import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Emu
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import TNR, _esc                                    # noqa: E402

HERE = os.path.dirname(os.path.abspath(__file__))
FIG_WIDTH = Emu(5486400)          # 6.0 in, matching the Part 1 figures
BM_ID = [7100]

SUBTITLE = ("Centralized Role-Based Event Coordination Platform — FYP Part 2 "
            "Development and Documentation Timeline")

FIGURES = [
    (os.path.join(HERE, "gantt2_part1.png"), "Gantt Chart (FYP Part 2) – part 1."),
    (os.path.join(HERE, "gantt2_part2.png"), "Gantt Chart (FYP Part 2) – part 2."),
]

# Task | Duration | Start | End | Status — appended after the final existing row
PLAN_ROWS = [
    ("9. Chapter 4 — Design and Implementation", "13 Days", "10 Aug 2026", "26 Aug 2026", "Completed"),
    ("9.1 Introduction", "1 Day", "10 Aug 2026", "10 Aug 2026", "Completed"),
    ("9.2 Design", "4 Days", "11 Aug 2026", "14 Aug 2026", "Completed"),
    ("9.3 Database Design", "2 Days", "17 Aug 2026", "18 Aug 2026", "Completed"),
    ("9.4 Interface Design", "2 Days", "19 Aug 2026", "20 Aug 2026", "Completed"),
    ("9.5 Implementation", "3 Days", "21 Aug 2026", "25 Aug 2026", "Completed"),
    ("9.6 Sample Codes", "1 Day", "26 Aug 2026", "26 Aug 2026", "Completed"),
    ("9.7 Summary", "1 Day", "26 Aug 2026", "26 Aug 2026", "Completed"),
    ("10. Chapter 5 — Results and Discussions", "5 Days", "27 Aug 2026", "02 Sep 2026", "Completed"),
    ("10.1 Introduction", "1 Day", "27 Aug 2026", "27 Aug 2026", "Completed"),
    ("10.2 Test Plan", "2 Days", "27 Aug 2026", "28 Aug 2026", "Completed"),
    ("10.3 Testing Results and Discussion", "3 Days", "31 Aug 2026", "02 Sep 2026", "Completed"),
    ("10.4 Summary", "1 Day", "02 Sep 2026", "02 Sep 2026", "Completed"),
    ("11. Chapter 6 — Conclusion", "3 Days", "03 Sep 2026", "07 Sep 2026", "Completed"),
    ("11.1 Critical Evaluation", "1 Day", "03 Sep 2026", "03 Sep 2026", "Completed"),
    ("11.2 Limitation", "1 Day", "04 Sep 2026", "04 Sep 2026", "Completed"),
    ("11.3 Recommendation", "1 Day", "07 Sep 2026", "07 Sep 2026", "Completed"),
    ("12. Appendices and Gantt Chart Update", "1 Day", "08 Sep 2026", "08 Sep 2026", "Completed"),
    ("13. Full Report Proofreading and Turnitin Check", "2 Days", "09 Sep 2026", "10 Sep 2026", "Completed"),
]


def caption_paragraph(doc, text: str):
    """`Figure <SEQ>: text`, centred and bold, exactly as the report's own."""
    rpr = f'<w:rPr>{TNR}<w:b/><w:bCs/></w:rPr>'
    ppr = (f'<w:pPr><w:pStyle w:val="Caption"/>'
           f'<w:spacing w:line="360" w:lineRule="auto"/>'
           f'<w:jc w:val="center"/>{rpr}</w:pPr>')
    body = (
        f'<w:r>{rpr}<w:t xml:space="preserve">Figure </w:t></w:r>'
        f'<w:r>{rpr}<w:fldChar w:fldCharType="begin"/></w:r>'
        f'<w:r>{rpr}<w:instrText xml:space="preserve"> SEQ Figure \\* ARABIC </w:instrText></w:r>'
        f'<w:r>{rpr}<w:fldChar w:fldCharType="separate"/></w:r>'
        f'<w:r><w:rPr>{TNR}<w:b/><w:bCs/><w:noProof/></w:rPr><w:t>1</w:t></w:r>'
        f'<w:r>{rpr}<w:fldChar w:fldCharType="end"/></w:r>'
        f'<w:r>{rpr}<w:t xml:space="preserve">: {_esc(text)}</w:t></w:r>')
    return parse_xml(f'<w:p {nsdecls("w")}>{ppr}{body}</w:p>')


def body_paragraph(model, text: str):
    """A plain body paragraph carrying the model paragraph's own formatting."""
    new = copy.deepcopy(model._p)
    for attr in (qn("w14:paraId"), qn("w14:textId")):
        new.attrib.pop(attr, None)
    ppr = new.find(qn("w:pPr"))
    for child in list(new):
        if child.tag != qn("w:pPr"):
            new.remove(child)
    run = new.makeelement(qn("w:r"), {})
    rpr = ppr.find(qn("w:rPr")) if ppr is not None else None
    if rpr is not None:
        run.append(copy.deepcopy(rpr))
    t = new.makeelement(qn("w:t"), {})
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    new.append(run)
    return new


def picture_paragraph(doc, path: str, width=None):
    """Build the picture in a scratch paragraph, then hand back its XML."""
    para = doc.add_paragraph()
    para.alignment = 1                                   # centred
    para.add_run().add_picture(path, width=width or FIG_WIDTH)
    element = para._p
    element.getparent().remove(element)
    return element


def add_plan_rows(table) -> int:
    last = table.rows[-1]._tr
    added = 0
    for values in PLAN_ROWS:
        row = copy.deepcopy(last)
        for attr in (qn("w14:paraId"), qn("w14:textId")):
            for node in row.iter():
                node.attrib.pop(attr, None)
        cells = row.findall(qn("w:tc"))
        for cell, value in zip(cells, values):
            texts = cell.findall(".//" + qn("w:t"))
            if not texts:
                continue
            texts[0].text = value
            texts[0].set(qn("xml:space"), "preserve")
            for extra in texts[1:]:
                extra.text = ""
        table._tbl.append(row)
        added += 1
    return added


def main() -> int:
    path = sys.argv[1]
    doc = docx.Document(path)

    # ── anchor: the caption of the last existing Gantt figure ───────────
    # The List of Figures repeats every caption, so the anchor must be the real
    # caption inside Appendix D: a Caption-styled paragraph after that heading.
    paras = doc.paragraphs
    appendix = next(i for i, p in enumerate(paras)
                    if p.style.name.startswith("Heading") and "Appendix D" in p.text)
    anchor = next(p for p in paras[appendix:]
                  if p.style.name == "Caption" and "Gantt Chart – part 2" in p.text)
    model = next(p for p in paras[appendix:] if "IR Semester Timeline" in p.text)

    cursor = anchor._p
    blocks = [body_paragraph(model, SUBTITLE)]
    for image, label in FIGURES:
        blocks.append(picture_paragraph(doc, image))
        blocks.append(caption_paragraph(doc, label))
    for block in blocks:
        cursor.addnext(block)
        cursor = block
    print(f"  inserted after '{anchor.text.strip()[:34]}': "
          f"1 subtitle + {len(FIGURES)} figures + {len(FIGURES)} captions")

    # ── project plan: append, never overwrite ───────────────────────────
    plan = None
    for table in doc.tables:
        head = [c.text.strip() for c in table.rows[0].cells]
        if head[:3] == ["Task", "Duration", "Start Date"]:
            plan = table
            break
    if plan is None:
        raise SystemExit("project plan table not found")
    before = len(plan.rows)
    added = add_plan_rows(plan)
    print(f"  project plan rows {before} -> {before + added}  (+{added}, none replaced)")

    doc.save(path)
    return 0


if __name__ == "__main__":
    sys.exit(main())
