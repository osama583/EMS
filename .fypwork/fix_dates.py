# -*- coding: utf-8 -*-
"""Put the FYP Part 2 timeline on the dates that actually applied.

The Part 2 work ran from 10 June 2026 to the submission on 2 September 2026,
not the April-to-September window carried over from the earlier draft. This
swaps the two regenerated Gantt images and rewrites every Part 2 row in the
project plan. The IR rows (1 to 34) and the IR Gantt are not touched.
"""
from __future__ import annotations

import os
import sys

import docx
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))

# caption fragment -> regenerated image
FIGURES = {
    "Gantt Chart (FYP Part 2) – part 1": os.path.join(HERE, "gantt2_part1.png"),
    "Gantt Chart (FYP Part 2) – part 2": os.path.join(HERE, "gantt2_part2.png"),
}

# Task label -> (duration, start, end). Only Part 2 rows.
DATES = {
    "5. FYP Part 2 — System Development": ("40 Days", "10 Jun 2026", "04 Aug 2026"),
    "5.1 System and Database Design": ("10 Days", "10 Jun 2026", "23 Jun 2026"),
    "5.2 Backend Development": ("15 Days", "24 Jun 2026", "14 Jul 2026"),
    "5.3 Frontend Development": ("15 Days", "15 Jul 2026", "04 Aug 2026"),
    "5.4 AI Assistant and Dashboards": ("5 Days", "05 Aug 2026", "11 Aug 2026"),
    "6. Testing and Evaluation": ("5 Days", "12 Aug 2026", "18 Aug 2026"),
    "7. Final Report Documentation": ("10 Days", "19 Aug 2026", "01 Sep 2026"),
    "8. Final Submission": ("1 Day", "02 Sep 2026", "02 Sep 2026"),

    "9. Chapter 4 — Design and Implementation": ("5 Days", "19 Aug 2026", "25 Aug 2026"),
    "9.1 Introduction": ("1 Day", "19 Aug 2026", "19 Aug 2026"),
    "9.2 Design": ("2 Days", "19 Aug 2026", "20 Aug 2026"),
    "9.3 Database Design": ("1 Day", "21 Aug 2026", "21 Aug 2026"),
    "9.4 Interface Design": ("1 Day", "24 Aug 2026", "24 Aug 2026"),
    "9.5 Implementation": ("2 Days", "24 Aug 2026", "25 Aug 2026"),
    "9.6 Sample Codes": ("1 Day", "25 Aug 2026", "25 Aug 2026"),
    "9.7 Summary": ("1 Day", "25 Aug 2026", "25 Aug 2026"),

    "10. Chapter 5 — Results and Discussions": ("2 Days", "26 Aug 2026", "27 Aug 2026"),
    "10.1 Introduction": ("1 Day", "26 Aug 2026", "26 Aug 2026"),
    "10.2 Test Plan": ("1 Day", "26 Aug 2026", "26 Aug 2026"),
    "10.3 Testing Results and Discussion": ("1 Day", "27 Aug 2026", "27 Aug 2026"),
    "10.4 Summary": ("1 Day", "27 Aug 2026", "27 Aug 2026"),

    "11. Chapter 6 — Conclusion": ("2 Days", "28 Aug 2026", "31 Aug 2026"),
    "11.1 Critical Evaluation": ("1 Day", "28 Aug 2026", "28 Aug 2026"),
    "11.2 Limitation": ("1 Day", "28 Aug 2026", "28 Aug 2026"),
    "11.3 Recommendation": ("1 Day", "31 Aug 2026", "31 Aug 2026"),

    "12. Appendices and Gantt Chart Update": ("1 Day", "01 Sep 2026", "01 Sep 2026"),
    "13. Full Report Proofreading and Turnitin Check": ("1 Day", "01 Sep 2026", "01 Sep 2026"),
}


def set_cell(cell, value: str) -> None:
    texts = cell._tc.findall(".//" + qn("w:t"))
    if not texts:
        cell.paragraphs[0].add_run(value)
        return
    texts[0].text = value
    texts[0].set(qn("xml:space"), "preserve")
    for extra in texts[1:]:
        extra.text = ""


def main() -> int:
    path = sys.argv[1]
    doc = docx.Document(path)
    body = list(doc.element.body)

    # ── the two Part 2 Gantt images ────────────────────────────────────
    swapped = 0
    for i, ch in enumerate(body):
        if ch.tag != qn("w:p"):
            continue
        # The List of Figures repeats every caption, so only a Caption-styled
        # paragraph counts; matching its index entry would rewrite whichever
        # image happens to sit above the list.
        style = ch.find(qn("w:pPr"))
        style = style.find(qn("w:pStyle")) if style is not None else None
        if style is None or style.get(qn("w:val")) != "Caption":
            continue
        text = "".join(t.text or "" for t in ch.iter(qn("w:t")))
        match = next((v for k, v in FIGURES.items() if k in text), None)
        if match is None:
            continue
        # the picture sits in the paragraph immediately above its caption
        for prev in reversed(body[:i]):
            blips = prev.findall(".//" + qn("a:blip"))
            if not blips:
                continue
            part = doc.part.related_parts[blips[0].get(qn("r:embed"))]
            before = len(part.blob)
            part._blob = open(match, "rb").read()
            print(f"  figure  {os.path.basename(match)}  {before:,} -> {len(part.blob):,} bytes")
            swapped += 1
            break
    if swapped != len(FIGURES):
        raise SystemExit(f"swapped {swapped} of {len(FIGURES)} figures")

    # ── project plan dates ─────────────────────────────────────────────
    plan = next(t for t in doc.tables
                if t.rows and [c.text.strip() for c in t.rows[0].cells][:2] == ["Task", "Duration"])
    updated, missing = 0, set(DATES)
    for row in plan.rows[1:]:
        label = row.cells[0].text.strip()
        if label not in DATES:
            continue
        duration, start, end = DATES[label]
        set_cell(row.cells[1], duration)
        set_cell(row.cells[2], start)
        set_cell(row.cells[3], end)
        set_cell(row.cells[4], "Completed")
        missing.discard(label)
        updated += 1
    print(f"  plan rows updated: {updated}")
    if missing:
        print(f"  NOT FOUND: {sorted(missing)}")

    doc.save(path)
    return 0


if __name__ == "__main__":
    sys.exit(main())
