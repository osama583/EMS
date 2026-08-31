"""Building blocks for appending chapters to the FYP report.

Everything here reproduces the formatting the document already uses rather than
inventing its own, so inserted material is indistinguishable from what was
written by hand: Times New Roman throughout, 1.5 line spacing (w:line 360),
justified body text, and captions that carry a real SEQ field so Word numbers
them rather than trusting a number typed into the text.
"""
from __future__ import annotations

import os

import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Emu
from PIL import Image

# --- page geometry ----------------------------------------------------------
EMU_PER_IN = 914400
EMU_PER_DXA = 635
TEXT_WIDTH = 5731510          # 6.268 in, measured from the document's own sectPr
PAIR_HEIGHT = 2700000         # both panes of a desktop/mobile pair share this
CODE_MAX_HEIGHT = 5760720     # 6.3 in - leaves room for the caption and the
                              # first paragraph, so a listing does not strand
                              # its own heading on the page above it

TNR = ('<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" '
       'w:cs="Times New Roman"/><w:color w:val="auto"/>')


def _esc(text: str) -> str:
    return (text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))


class Builder:
    """Appends blocks to the end of the body and remembers what it added.

    Working at the end and relocating afterwards keeps the XML simple: python-docx
    handles image relationships and table grids correctly only for content it
    appends itself, so nothing here has to hand-roll a relationship id.
    """

    def __init__(self, doc: docx.Document):
        self.doc = doc
        self.body = doc.element.body
        self.added: list = []

    # -- internals -----------------------------------------------------------
    def _append(self, element):
        self.body.append(element)
        self.added.append(element)
        return element

    @staticmethod
    def _set_tblPr(tbl, xml: str):
        existing = tbl.find(qn('w:tblPr'))
        new = parse_xml(xml)
        if existing is not None:
            tbl.replace(existing, new)
        else:
            tbl.insert(0, new)

    def _p(self, inner_ppr: str, runs: str):
        xml = f'<w:p {nsdecls("w")}>{inner_ppr}{runs}</w:p>'
        return self._append(parse_xml(xml))

    # -- text ----------------------------------------------------------------
    def heading(self, text: str, level: int):
        """A numbered heading in the document's own house style."""
        size = {1: '', 2: '<w:sz w:val="32"/><w:szCs w:val="32"/>',
                3: '<w:sz w:val="28"/><w:szCs w:val="28"/>', 4: ''}[level]
        bold = '<w:b/><w:bCs/>' if level == 4 else ''
        indent = '<w:ind w:left="720" w:hanging="720"/>' if level == 1 else ''
        rpr = f'<w:rPr>{TNR}{bold}{size}</w:rPr>'
        ppr = (f'<w:pPr><w:pStyle w:val="Heading{level}"/>'
               f'<w:keepNext/><w:keepLines/>'
               f'<w:spacing w:line="360" w:lineRule="auto"/>{indent}'
               f'<w:jc w:val="both"/>{rpr}</w:pPr>')
        return self._p(ppr, f'<w:r>{rpr}<w:t xml:space="preserve">{_esc(text)}</w:t></w:r>')

    def body_text(self, text: str, indent_first: bool = True,
                  keep_next: bool = False):
        """A justified body paragraph. Body paragraphs in this report open with a
        literal tab rather than a style indent, so that is reproduced here."""
        rpr = f'<w:rPr>{TNR}</w:rPr>'
        keep = '<w:keepNext/>' if keep_next else ''
        ppr = (f'<w:pPr>{keep}<w:spacing w:line="360" w:lineRule="auto"/>'
               f'<w:jc w:val="both"/>{rpr}</w:pPr>')
        lead = '\t' if indent_first else ''
        return self._p(ppr, f'<w:r>{rpr}<w:t xml:space="preserve">{_esc(lead + text)}</w:t></w:r>')

    def plain(self, text: str = '', centre: bool = False, compact: bool = False):
        """`compact` drops to single spacing with no space after, which is what
        the short rating-scale lines want - at 1.5 they take half a page."""
        rpr = f'<w:rPr>{TNR}</w:rPr>'
        jc = '<w:jc w:val="center"/>' if centre else ''
        spacing = ('<w:spacing w:line="240" w:lineRule="auto" w:after="0"/>'
                   if compact else '<w:spacing w:line="360" w:lineRule="auto"/>')
        ppr = f'<w:pPr>{spacing}{jc}{rpr}</w:pPr>'
        run = f'<w:r>{rpr}<w:t xml:space="preserve">{_esc(text)}</w:t></w:r>' if text else ''
        return self._p(ppr, run)

    def page_break(self):
        return self._p('', f'<w:r><w:br w:type="page"/></w:r>')

    # -- captions ------------------------------------------------------------
    def caption(self, kind: str, text: str, bookmark: str | None = None,
                bookmark_id: int = 0):
        """`Figure <SEQ>: text`, where <SEQ> is a live field.

        A caption typed as literal text is why the existing chapters disagree with
        their own List of Figures. Every caption produced here numbers itself, so
        inserting a figure renumbers everything after it on the next field update.
        The bookmark is what a cross-reference in the prose points at.
        """
        rpr = f'<w:rPr>{TNR}<w:b/><w:bCs/></w:rPr>'
        bm_start = (f'<w:bookmarkStart w:id="{bookmark_id}" w:name="{bookmark}"/>'
                    if bookmark else '')
        bm_end = f'<w:bookmarkEnd w:id="{bookmark_id}"/>' if bookmark else ''
        runs = (
            f'{bm_start}'
            f'<w:r>{rpr}<w:t xml:space="preserve">{kind} </w:t></w:r>'
            f'<w:r>{rpr}<w:fldChar w:fldCharType="begin"/></w:r>'
            f'<w:r>{rpr}<w:instrText xml:space="preserve"> SEQ {kind} \\* ARABIC </w:instrText></w:r>'
            f'<w:r>{rpr}<w:fldChar w:fldCharType="separate"/></w:r>'
            f'<w:r><w:rPr>{TNR}<w:b/><w:bCs/><w:noProof/></w:rPr><w:t>1</w:t></w:r>'
            f'<w:r>{rpr}<w:fldChar w:fldCharType="end"/></w:r>'
            f'<w:r>{rpr}<w:t xml:space="preserve">: {_esc(text)}</w:t></w:r>'
            f'{bm_end}'
        )
        ppr = (f'<w:pPr><w:pStyle w:val="Caption"/>'
               f'<w:spacing w:line="360" w:lineRule="auto"/>'
               f'<w:jc w:val="center"/>{rpr}</w:pPr>')
        return self._p(ppr, runs)

    # -- figures -------------------------------------------------------------
    def _picture_cell(self, cell, path: str, width: int, height: int):
        cell_p = cell.paragraphs[0]
        cell_p.paragraph_format.space_after = 0
        cell_p.paragraph_format.keep_with_next = True
        cell_p._p.get_or_add_pPr().append(
            parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))
        cell_p.add_run().add_picture(path, width=Emu(width), height=Emu(height))

    def figure_pair(self, desktop: str, mobile: str, height: int = PAIR_HEIGHT):
        """The approved two-pane layout: desktop and mobile at equal height, side
        by side, borderless and centred, with one caption underneath."""
        dw, dh = Image.open(desktop).size
        mw, mh = Image.open(mobile).size
        d_w = round(height * dw / dh)
        m_w = round(height * mw / mh)

        over = (d_w + m_w) - TEXT_WIDTH
        if over > 0:                       # never let a figure run into the margin
            scale = TEXT_WIDTH / (d_w + m_w)
            height = round(height * scale)
            d_w, m_w = round(d_w * scale), round(m_w * scale)

        table = self.doc.add_table(rows=1, cols=2)
        self.added.append(table._tbl)
        d_col = round(d_w / EMU_PER_DXA)
        m_col = 9026 - d_col
        self._set_tblPr(table._tbl, (
            f'<w:tblPr {nsdecls("w")}>'
            f'<w:tblW w:type="auto" w:w="0"/><w:tblLayout w:type="fixed"/>'
            f'<w:tblBorders><w:top w:val="none" w:sz="0" w:space="0"/>'
            f'<w:left w:val="none" w:sz="0" w:space="0"/>'
            f'<w:bottom w:val="none" w:sz="0" w:space="0"/>'
            f'<w:right w:val="none" w:sz="0" w:space="0"/>'
            f'<w:insideH w:val="none" w:sz="0" w:space="0"/>'
            f'<w:insideV w:val="none" w:sz="0" w:space="0"/></w:tblBorders>'
            f'<w:tblCellMar>'
            f'<w:left w:w="0" w:type="dxa"/><w:right w:w="0" w:type="dxa"/>'
            f'</w:tblCellMar><w:jc w:val="center"/></w:tblPr>'))
        grid = table._tbl.find(qn('w:tblGrid'))
        for col, w in zip(grid.findall(qn('w:gridCol')), (d_col, m_col)):
            col.set(qn('w:w'), str(w))
        for cell, w in zip(table.rows[0].cells, (d_col, m_col)):
            cell.width = Emu(w * EMU_PER_DXA)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        self._picture_cell(table.rows[0].cells[0], desktop, d_w, height)
        self._picture_cell(table.rows[0].cells[1], mobile, m_w, height)
        return table

    def figure_single(self, path: str, max_width: int = TEXT_WIDTH,
                      max_height: int = CODE_MAX_HEIGHT):
        w, h = Image.open(path).size
        width = max_width
        height = round(width * h / w)
        if height > max_height:
            height = max_height
            width = round(height * w / h)
        rpr = f'<w:rPr>{TNR}</w:rPr>'
        ppr = (f'<w:pPr><w:keepNext/>'
               f'<w:spacing w:line="240" w:lineRule="auto" w:after="0"/>'
               f'<w:jc w:val="center"/>{rpr}</w:pPr>')
        para = self._p(ppr, '')
        from docx.text.paragraph import Paragraph
        Paragraph(para, self.doc).add_run().add_picture(
            path, width=Emu(width), height=Emu(height))
        return para

    # -- tables --------------------------------------------------------------
    def data_table(self, rows: list[list[str]], widths: list[int] | None = None,
                   header: bool = True, size: int = 18):
        """A bordered data table sized to the text column."""
        cols = len(rows[0])
        table = self.doc.add_table(rows=len(rows), cols=cols)
        self.added.append(table._tbl)
        self._set_tblPr(table._tbl, (
            f'<w:tblPr {nsdecls("w")}>'
            f'<w:tblW w:type="dxa" w:w="9026"/><w:tblLayout w:type="fixed"/>'
            f'<w:tblBorders>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'</w:tblBorders><w:jc w:val="center"/></w:tblPr>'))
        if widths is None:
            widths = [9026 // cols] * cols
            widths[-1] += 9026 - sum(widths)
        grid = table._tbl.find(qn('w:tblGrid'))
        for col, w in zip(grid.findall(qn('w:gridCol')), widths):
            col.set(qn('w:w'), str(w))
        for r, row in enumerate(rows):
            tr = table.rows[r]
            if header and r == 0:
                tr._tr.get_or_add_trPr().append(
                    parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
            for c, value in enumerate(row):
                cell = tr.cells[c]
                cell.width = Emu(widths[c] * EMU_PER_DXA)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if header and r == 0:
                    cell._tc.get_or_add_tcPr().append(parse_xml(
                        f'<w:shd {nsdecls("w")} w:val="clear" w:fill="D9E2F3"/>'))
                p = cell.paragraphs[0]
                bold = '<w:b/><w:bCs/>' if (header and r == 0) else ''
                rpr = (f'<w:rPr>{TNR}{bold}<w:sz w:val="{size}"/>'
                       f'<w:szCs w:val="{size}"/></w:rPr>')
                p._p.append(parse_xml(
                    f'<w:pPr {nsdecls("w")}><w:spacing w:line="240" '
                    f'w:lineRule="auto" w:after="40"/>{rpr}</w:pPr>'))
                p._p.append(parse_xml(
                    f'<w:r {nsdecls("w")}>{rpr}'
                    f'<w:t xml:space="preserve">{_esc(str(value))}</w:t></w:r>'))
        return table

    # -- placement -----------------------------------------------------------
    def move_before(self, anchor):
        """Relocate everything built so far to sit immediately before `anchor`."""
        for element in self.added:
            anchor.addprevious(element)
        self.added = []
