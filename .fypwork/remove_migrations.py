# -*- coding: utf-8 -*-
"""Take the migration mechanism out of the report.

The migrations directory is how the schema reached its current state, not a
feature of the delivered system - the same category as the .env file. It is
therefore removed from the documentation rather than described in it.

Four consequences are handled here, not just the wording:

  * the ERD figure is replaced with one regenerated without the ledger table;
  * the data dictionary row for schema_migrations is replaced by
    club_membership_log, a real product table that was missing from it, so the
    dictionary still documents sixty-eight tables and all of them are ours;
  * the foreign key count rises from 111 to 114, because the stale snapshot the
    old figure was drawn from predated club_membership_log;
  * a cross-reference in 4.2.2 that pointed at 2.4.6 is corrected to 2.4.4.
"""
from __future__ import annotations

import sys

import docx
from docx.oxml.ns import qn

ERD_PNG = r"C:\Users\natsu\Desktop\Osama\digrams\png\00-erd.png"

# (find, replace) applied to whole paragraphs, matched on a unique fragment
REWRITES = [
    # 2.4.4 justification - drop the migration runner, keep the point about
    # what Supabase does and does not own
    ("while the schema stays under the project's own versioned migration runner.",
     "while the database itself remains a standard PostgreSQL instance that could be moved to "
     "another host without rewriting the application."),

    # 4.2.2 - remove the migration sentence, and correct the cross-reference,
    # which pointed at the deployment section rather than the database one
    ("Its schema is applied by an ordered set of migration files rather than by manual changes, "
     "which means the database can be rebuilt from an empty state into a known configuration. "
     "This is the arrangement described and justified in Section 2.4.6 of this report,",
     "The application therefore holds no database of its own, and every environment reads and "
     "writes the same authoritative copy. This is the arrangement described and justified in "
     "Section 2.4.4 of this report,"),

    # 4.3 - drop the migration clause and correct the foreign key count
    ("sixty-eight tables joined by one hundred and eleven foreign key relationships, applied "
     "through an ordered sequence of migration files rather than by manual alteration, and "
     "hosted as a managed PostgreSQL instance.",
     "sixty-eight tables joined by one hundred and fourteen foreign key relationships, hosted "
     "as a managed PostgreSQL instance."),

    # 4.3.1 - keep the introspection claim, drop the migration illustration
    ("a table added by a later migration appears in the diagram automatically rather than being "
     "forgotten.",
     "every table, column, key and relationship shown in it is read from the database itself."),
]

# whole paragraphs to delete
DELETIONS = [
    "Schema changes applied through a versioned migration runner",
]

# data dictionary: the ledger row gives way to a product table that was missing
ROW_SWAP = (
    "schema_migrations",
    ["club_membership_log",
     "Audit trail of club membership transitions, recording joining, leaving, removal and "
     "president handover, together with the acting user where one was recorded.",
     "club_membership_log_id, club_id, subject_user_id, actor_user_id, action, role_label, "
     "occurred_at"],
)


def set_runs(p, text: str) -> bool:
    """Write text into the paragraph's first text run, clear the others."""
    runs = [n for r in p._p.findall(qn("w:r")) for n in r]
    first = None
    for node in runs:
        if node.tag == qn("w:t"):
            if first is None:
                node.text = text
                node.set(qn("xml:space"), "preserve")
                first = node
            else:
                node.text = ""
    return first is not None


def set_cell(cell, text: str) -> None:
    para = cell.paragraphs[0]
    runs = para._p.findall(qn("w:r"))
    if not runs:
        para.add_run(text)
        return
    first = runs[0].find(qn("w:t"))
    if first is None:
        first = runs[0].makeelement(qn("w:t"), {})
        runs[0].append(first)
    first.text = text
    first.set(qn("xml:space"), "preserve")
    for extra in runs[1:]:
        for t in extra.findall(qn("w:t")):
            t.text = ""


def main() -> int:
    path = sys.argv[1]
    doc = docx.Document(path)

    # --- prose -------------------------------------------------------------
    done = 0
    for old, new in REWRITES:
        hits = [p for p in doc.paragraphs if old in p.text]
        if len(hits) != 1:
            raise SystemExit(f"matched {len(hits)} paragraphs for: {old[:60]!r}")
        set_runs(hits[0], hits[0].text.replace(old, new))
        done += 1
        print(f"  rewrote  {old[:56]}")

    removed = 0
    for needle in DELETIONS:
        hits = [p for p in doc.paragraphs if needle in p.text]
        if len(hits) != 1:
            raise SystemExit(f"matched {len(hits)} paragraphs for: {needle!r}")
        hits[0]._p.getparent().remove(hits[0]._p)
        removed += 1
        print(f"  deleted  {needle[:56]}")

    # --- data dictionary row ----------------------------------------------
    target, values = ROW_SWAP
    swapped = 0
    for table in doc.tables:
        for row in table.rows:
            if row.cells[0].text.strip() == target:
                for cell, value in zip(row.cells, values):
                    set_cell(cell, value)
                swapped += 1
    if swapped != 1:
        raise SystemExit(f"data dictionary row swap hit {swapped} rows")
    print(f"  swapped  data dictionary row {target} -> {values[0]}")

    # --- ERD figure --------------------------------------------------------
    with open(ERD_PNG, "rb") as fh:
        blob = fh.read()
    part = next(p for p in doc.part.package.parts
                if str(p.partname) == "/word/media/image45.png")
    before = len(part.blob)
    part._blob = blob
    print(f"  figure   image45.png {before:,} -> {len(blob):,} bytes")

    doc.save(path)

    check = docx.Document(path)
    left = [p.text.strip()[:70] for p in check.paragraphs if "migrat" in p.text.lower()]
    left += [c.text.strip()[:70] for t in check.tables for r in t.rows for c in r.cells
             if "migrat" in c.text.lower()]
    print(f"\n  paragraphs rewritten : {done}")
    print(f"  paragraphs deleted   : {removed}")
    print(f"  remaining 'migrat'   : {left or 'none'}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
