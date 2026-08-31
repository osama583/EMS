import docx, sys
from docx.oxml.ns import qn
p = r"C:\Users\natsu\Downloads\Osamah_Ahmed_Moahmmed_Al-Naggar-TP078781-APD3F2601CS-FYP_2 - Copy.docx"
d = docx.Document(p)
body = d.element.body
print("total block children:", len(body))
paras = d.paragraphs
print("total paragraphs:", len(paras))
print("total tables:", len(d.tables))
print("sections:", len(d.sections))
print("=== styles used (heading-ish) ===")
from collections import Counter
c = Counter(x.style.name for x in paras)
for k,v in sorted(c.items(), key=lambda t:-t[1]):
    print(f"  {v:5d}  {k}")
