"""Edit the finished report in place: the UAT section and the acknowledgement.

Nothing else in the document is touched. Every other paragraph, table, figure
and field is left exactly as it was, so this is a surgical edit rather than a
regeneration.
"""
from __future__ import annotations

import sys

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.path.insert(0, '.')
import content_ch5_uat as UAT      # noqa: E402


def set_cell(cell, text: str):
    """Replace a cell's text, keeping the run formatting already in place."""
    para = cell.paragraphs[0]
    runs = para._p.findall(qn('w:r'))
    if not runs:
        return
    first = runs[0].find(qn('w:t'))
    if first is None:
        return
    first.text = text
    first.set(qn('xml:space'), 'preserve')
    for extra in runs[1:]:
        t = extra.find(qn('w:t'))
        if t is not None:
            t.text = ''


def set_caption_label(paragraph, label: str):
    """A caption is `Table <SEQ>: <label>`; only the trailing run is the label."""
    runs = paragraph._p.findall(qn('w:r'))
    for run in reversed(runs):
        t = run.find(qn('w:t'))
        if t is not None and t.text and t.text.startswith(':'):
            t.text = f": {label}"
            t.set(qn('xml:space'), 'preserve')
            return True
    return False


def drop_age_row(table) -> bool:
    for row in table.rows:
        if row.cells[0].text.strip() == "Age":
            row._tr.getparent().remove(row._tr)
            return True
    return False


def replace_paragraph_text(paragraph, text: str):
    runs = paragraph._p.findall(qn('w:r'))
    if not runs:
        return False
    first = runs[0].find(qn('w:t'))
    if first is None:
        return False
    lead = '\t' if (first.text or '').startswith('\t') else ''
    first.text = lead + text
    first.set(qn('xml:space'), 'preserve')
    for extra in runs[1:]:
        t = extra.find(qn('w:t'))
        if t is not None:
            t.text = ''
    return True


def main() -> int:
    path = sys.argv[1]
    doc = docx.Document(path)
    body = list(doc.element.body)

    # --- acknowledgement ---------------------------------------------------
    ack = 0
    for para in doc.paragraphs:
        if "Adel" in para.text and "friends" in para.text:
            for run in para._p.findall(qn('w:r')):
                t = run.find(qn('w:t'))
                if t is None or not t.text:
                    continue
                if t.text == ' and ':
                    t.text = ', '
                    t.set(qn('xml:space'), 'preserve')
                elif t.text == 'Sohaib':
                    t.text = 'Sohaib and Ibrahim'
            ack += 1
            break

    # --- UAT prose ---------------------------------------------------------
    prose = 0
    swaps = [
        ("Five testers were recruited", UAT.UAT_PLAN[1]),
        ("Each participant was given the running system", UAT.UAT_PLAN[2]),
    ]
    old_discussion_keys = [
        "The responses are read together rather than averaged",
        "Across the twenty interface ratings",
        "The functionality criteria are the more significant",
        "The comments converge on a theme",
        "Taken with the unit testing results",
    ]
    new_discussion = list(UAT.UAT_DISCUSSION)
    for para in doc.paragraphs:
        text = para.text.strip()
        for key, replacement in swaps:
            if text.startswith(key):
                replace_paragraph_text(para, replacement)
                prose += 1
        for i, key in enumerate(old_discussion_keys):
            if text.startswith(key):
                # keep positional order; the paired-departments paragraph is new
                idx = i if i < 3 else i + 1
                replace_paragraph_text(para, new_discussion[idx])
                prose += 1

    # the tasks paragraph gains a follow-on paragraph about the paired departments
    inserted = 0
    for para in doc.paragraphs:
        if para.text.strip().startswith(UAT.UAT_PLAN[2][:45]):
            clone = para._p.makeelement(qn('w:p'), {})
            for child in para._p:
                import copy
                clone.append(copy.deepcopy(child))
            para._p.addnext(clone)
            replace_paragraph_text(Paragraph(clone, doc), UAT.UAT_PLAN[3])
            inserted += 1
            break

    # the discussion gains the paired-departments paragraph
    for para in doc.paragraphs:
        if para.text.strip().startswith(new_discussion[2][:45]):
            clone = para._p.makeelement(qn('w:p'), {})
            for child in para._p:
                import copy
                clone.append(copy.deepcopy(child))
            para._p.addnext(clone)
            replace_paragraph_text(Paragraph(clone, doc), new_discussion[3])
            inserted += 1
            break

    # --- UAT tables --------------------------------------------------------
    body = list(doc.element.body)
    profiles = []           # (index, table)
    for i, ch in enumerate(body):
        if ch.tag != qn('w:tbl'):
            continue
        table = Table(ch, doc)
        if table.rows and table.rows[0].cells[0].text.strip() == "Tester demographic profile":
            profiles.append((i, table))

    ages = names = roles = captions = comments = 0
    for n, (index, table) in enumerate(profiles):
        if drop_age_row(table):
            ages += 1
        if n == 0:                       # the blank instrument in 5.2.2
            continue
        tester = UAT.TESTERS[n - 1]
        for row in table.rows:
            key = row.cells[0].text.strip()
            if key == "Name":
                set_cell(row.cells[1], tester["name"]); names += 1
            elif key == "Role in the system":
                set_cell(row.cells[1], tester["role"]); roles += 1

        # the four captions belonging to this tester follow the profile table
        titles = ["User Acceptance Testing", "Interface Criteria Ratings",
                  "Functionality Criteria Responses", "Tester Comment"]
        seen = 0
        for ch in body[index:]:
            if ch.tag != qn('w:p'):
                if ch.tag == qn('w:tbl'):
                    tbl = Table(ch, doc)
                    if tbl.rows and tbl.rows[0].cells[0].text.strip() == "Tester comment:":
                        set_cell(tbl.rows[1].cells[0], tester["comment"])
                        comments += 1
                continue
            par = Paragraph(ch, doc)
            if par.style.name != 'Caption':
                continue
            if set_caption_label(par, f"{titles[seen]} — {tester['label']}"):
                captions += 1
            seen += 1
            if seen == 4:
                break

    doc.save(path)
    print(f"  acknowledgement updated : {ack}")
    print(f"  prose paragraphs replaced: {prose}")
    print(f"  paragraphs inserted      : {inserted}")
    print(f"  age rows removed         : {ages}")
    print(f"  names set                : {names}")
    print(f"  roles set                : {roles}")
    print(f"  captions relabelled      : {captions}")
    print(f"  comments replaced        : {comments}")
    return 0


sys.exit(main())
