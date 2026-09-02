# -*- coding: utf-8 -*-
"""Rebuild the appendices as A to H, matching the guidelines.

The report had six appendices at three different heading levels: A was Heading 1
(a sibling of APPENDICES rather than a child of it), E was Heading 3, and the
rest were Heading 2, so the table of contents nested them wrongly. Two required
appendices were missing entirely - the poster and the sample code.

    A  Project Proposal Form (PPF)          unchanged
    B  Ethics Forms                         unchanged
    C  Log Sheets                           unchanged
    D  Poster                               NEW
    E  Gantt Chart                          was D
    F  Sample Code Implementation           NEW
    G  Respondent Demographic Profile       was E
    H  Turnitin Similarity Report           was F

Renames run last-to-first so a new label never collides with one still in use.
"""
from __future__ import annotations

import os
import sys

import docx
from docx.oxml.ns import qn
from docx.shared import Emu

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from add_gantt2_and_plan import body_paragraph, caption_paragraph, picture_paragraph  # noqa: E402

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
WIDTH = Emu(5731510)                     # 6.27 in, the full text width

POSTER = os.path.join(ROOT, "poster-pack", "poster.png")
CODE_DIR = os.path.join(HERE, "appendix_code")

# renames, applied in this order
RELABEL = [
    ("Appendix F: Turnitin Similarity Report", "Appendix H: Turnitin Similarity Report"),
    ("Appendix E: Respondent Demographic Profile", "Appendix G: Respondent Demographic Profile"),
    ("Appendix D: Gantt Chart", "Appendix E: Gantt Chart"),
]

POSTER_CAPTION = "Project poster: the platform, its coordination workflow and its outcomes"

GITHUB = "https://github.com/osama583/EMS"

INTRO = [
    f"The complete source code for the delivered platform is published at {GITHUB}.",
    "The listings below present the security implementation of the system as a single subject "
    "rather than as unrelated samples. They follow the path a request takes through the server: "
    "the credential that opens a session, the tokens that carry it, the identity and page "
    "permissions that bound it, the query layer that binds every value it supplies, and the "
    "headers and error envelope that close it. The twelve listings in Section 4.6 show how a "
    "request is handled; these show what protects it.",
]

# (image stem, figure caption)
LISTINGS = [
    ("F_1", "Password hashing with a configurable bcrypt cost factor"),
    ("F_2", "Access and refresh tokens separated by a verified type claim"),
    ("F_3", "Authenticating a request and loading its roles from the database"),
    ("F_4", "Authority held as a role within an organisational unit"),
    ("F_5", "Resolving which pages a caller may open, from the grant tables"),
    ("F_6", "Pooled database access with every value bound as a query parameter"),
    ("F_7", "Cross-origin access restricted to an explicit origin list"),
    ("F_8", "Security headers applied uniformly to every response"),
    ("F_9", "Returning a user-safe error while logging the internal detail"),
    ("F_10", "Password reset without disclosing whether an account exists"),
]


def find_heading(doc, text: str):
    hits = [p for p in doc.paragraphs
            if p.style.name.startswith("Heading") and p.text.strip().startswith(text)]
    if len(hits) != 1:
        raise SystemExit(f"heading matched {len(hits)}: {text!r}")
    return hits[0]


def set_heading_text(para, text: str) -> None:
    runs = para._p.findall(qn("w:r"))
    first = None
    for run in runs:
        for t in run.findall(qn("w:t")):
            if first is None:
                t.text = text
                t.set(qn("xml:space"), "preserve")
                first = t
            else:
                t.text = ""
    if first is None:
        para.add_run(text)


def heading_like(doc, model, text: str):
    """A new appendix heading carrying the model heading's own formatting."""
    new = body_paragraph(model, text)
    return new


def main() -> int:
    path = sys.argv[1]
    doc = docx.Document(path)
    h2 = doc.styles["Heading 2"]

    # ── 1. relabel, last to first ──────────────────────────────────────
    for old, new in RELABEL:
        set_heading_text(find_heading(doc, old), new)
        print(f"  relabelled  {old:46s} -> {new}")

    # ── 2. one heading level for all of them ───────────────────────────
    levelled = 0
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading") and para.text.strip().startswith("Appendix "):
            if para.style.name != "Heading 2":
                para.style = h2
                levelled += 1
    print(f"  heading levels normalised to Heading 2: {levelled}")

    model = find_heading(doc, "Appendix B")          # heading formatting donor
    # Body text needs a body donor: cloning the heading would make the intro
    # paragraphs Heading 2 and put them in the table of contents.
    body_model = next(p for p in doc.paragraphs
                      if p.style.name == "Normal" and "IR Semester Timeline" in p.text)

    # ── 3. Appendix D: Poster, before the Gantt appendix ───────────────
    gantt = find_heading(doc, "Appendix E: Gantt Chart")
    blocks = [
        heading_like(doc, model, "Appendix D: Poster"),
        picture_paragraph(doc, POSTER, width=WIDTH),
        caption_paragraph(doc, POSTER_CAPTION),
    ]
    cursor = None
    for block in blocks:
        if cursor is None:
            gantt._p.addprevious(block)
        else:
            cursor.addnext(block)
        cursor = block
    print("  inserted    Appendix D: Poster (1 figure)")

    for para in doc.paragraphs:
        if para.text.strip() == "Appendix D: Poster":
            para.style = h2

    # ── 4. Appendix F: Sample Code, before Appendix G ──────────────────
    demographic = find_heading(doc, "Appendix G: Respondent Demographic Profile")
    blocks = [heading_like(doc, model,
                           "Appendix F: Sample Code Implementation (GitHub Link For Source Code)")]
    for text in INTRO:
        blocks.append(body_paragraph(body_model, text))
    for stem, caption in LISTINGS:
        blocks.append(picture_paragraph(doc, os.path.join(CODE_DIR, f"{stem}.png"), width=WIDTH))
        blocks.append(caption_paragraph(doc, caption))

    cursor = None
    for block in blocks:
        if cursor is None:
            demographic._p.addprevious(block)
        else:
            cursor.addnext(block)
        cursor = block
    for para in doc.paragraphs:
        if para.text.strip().startswith("Appendix F: Sample Code"):
            para.style = h2
    print(f"  inserted    Appendix F: Sample Code ({len(LISTINGS)} figures, "
          f"{len(INTRO)} paragraphs)")

    doc.save(path)

    order = [p.text.strip() for p in docx.Document(path).paragraphs
             if p.style.name.startswith("Heading") and p.text.strip().startswith("Appendix ")]
    print("\n  final order:")
    for line in order:
        print(f"    {line}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
