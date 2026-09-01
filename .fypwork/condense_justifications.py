# -*- coding: utf-8 -*-
"""Cut every Justification down to 1-3 paragraphs of 4 to 7 lines each.

At this page width a full body line runs to about fourteen words, so the
budget is roughly 56 to 98 words per paragraph. Each block keeps the same
three-part argument the lecturer asked for - why this, what else was available,
why the alternatives lost - but says it in a third of the words.

A block is a list of anchors naming the paragraphs it currently occupies and a
list of replacement texts. Paragraphs are rewritten pairwise; any anchor left
over is deleted, and any text left over is inserted after the last one. The
document is edited in place, so nothing else moves.
"""
from __future__ import annotations

import copy
import sys

import docx
from docx.oxml.ns import qn

# (anchors currently in the document, replacement paragraphs)
BLOCKS = [

    # ---- 2.4 opening -------------------------------------------------------
    ([
        "identifies and justifies all technical components",
    ], [
        "This section identifies and justifies all technical components selected for the "
        "platform. Each is presented in the same form: the alternatives that were realistically "
        "available, the reason each was set aside, and why the selected option fits this system. "
        "A list of advantages is not a justification, since the rejected options have advantages "
        "of their own, and several of them are the more common industry choice.",
    ]),

    # ---- 2.4.2 IDE ---------------------------------------------------------
    ([
        "ranked as the most popular developer environment tool",
        "Three alternatives were considered against it",
        "The deciding factor is the shape of the work",
    ], [
        "Justification: Tan et al. (2023) report that Visual Studio Code was the most used "
        "developer environment in the Stack Overflow 2021 survey, with 71.06% of over 80,000 "
        "respondents, and attribute this to its cross-platform support and its extension "
        "ecosystem. That matters here because development runs simultaneously across Angular in "
        "TypeScript and SCSS and Flask in Python, which one editor has to cover.",

        "Three alternatives were rejected. The JetBrains tools are the strongest per language but "
        "are split across two paid products, PyCharm for the backend and WebStorm for the "
        "frontend, meaning two editors for one codebase. Visual Studio is weighted towards the "
        ".NET ecosystem, which this project does not use. Sublime Text and Vim are free and fast "
        "but require language support and debugging to be assembled from plugins by hand.",

        "The deciding factor is the shape of the work rather than the depth of any one tool. A "
        "single developer moves between an Angular component and the Flask route that serves it "
        "many times in a session, so the cost that matters is crossing between four languages, "
        "not mastery of one. Visual Studio Code is the only candidate that covers all four free "
        "of charge with an integrated debugger.",
    ]),

    # ---- 2.4.3 A: Angular --------------------------------------------------
    ([
        "As stated by Vashisht (2021), Angular is distinguished",
        "Four alternatives were weighed against Angular",
        "Angular was selected because the properties this system needs",
    ], [
        "Vashisht (2021) identifies Angular's strength as managing highly complex applications "
        "through a strictly structured and modular platform. This is decisive here, because "
        "twelve user roles must be supported within one application: modularity lets the views "
        "for applicants, Heads of School and Department, the CFO, cafeteria users, club "
        "administrators and system administrators stay independent while sharing a codebase, with "
        "the client displaying only what the server permits.",

        "Four alternatives were rejected. React is a library rather than a framework, so routing, "
        "forms, HTTP and state become separate decisions the developer must research before "
        "writing a feature. Vue is less opinionated about structure and treats TypeScript as "
        "optional. Svelte compiles to smaller output, which Vashisht (2021) confirms, but the "
        "same comparison gives Angular the advantage in tooling for complex applications. Plain "
        "JavaScript would mean hand-writing routing and session handling.",

        "Angular was chosen because what this system needs, it supplies as framework features "
        "rather than as choices. Route guards, reactive forms, dependency injection and HTTP "
        "interceptors each have one documented mechanism. The guard and token interceptor in "
        "Sections 4.6.3 and 4.6.4 are configurations of those features; in React the same "
        "behaviour would be the developer's own code, and the access model would depend on it "
        "being written correctly in every place it was needed.",
    ]),

    # ---- 2.4.3 B: Flask ----------------------------------------------------
    ([
        "Flask is suitable for building RESTful APIs",
        "In the developed system, the backend is organised into resource-based modules",
        "Four backend alternatives were considered",
        "Flask was selected precisely because it imposes little",
    ], [
        "Flask provides a lightweight and flexible framework for the REST API this system needs, "
        "supporting authentication, database integration and modular endpoints (Albesher & "
        "Alfayez, 2024). In the delivered system the backend is organised into resource-based "
        "modules under one versioned API path, with the approval rules held in a service layer "
        "rather than in the request handlers, and every change to a request written inside a "
        "single transaction.",

        "Four alternatives were rejected. Django supplies an object-relational mapper, an admin "
        "interface and a permission system, but its permissions are per model and global per "
        "user, whereas authority here is a role held within an organisational unit, so Django's "
        "model would have been overridden rather than used. FastAPI's asynchronous handling does "
        "not help a database-bound approval queue. Node.js with Express would move the assistant "
        "and the scheduled jobs away from Python. Spring Boot and ASP.NET Core are "
        "disproportionate for one developer.",

        "Flask was chosen because it imposes little. The three-layer separation described in "
        "Section 4.2.2, where a route may not issue SQL and every change passes through the "
        "workflow service that writes its audit record, was designed around this workflow rather "
        "than inherited from a framework. A fuller framework would supply its own answers to "
        "authorisation, data access and routing, and this system needs a different answer in all "
        "three.",
    ]),

    # ---- 2.4.3 C: supporting libraries -------------------------------------
    ([
        "Two of the decisions represented in the table above",
        "The second is session handling",
    ], [
        "Two decisions in this table were made against real alternatives. SQLAlchemy is the usual "
        "pairing with Flask, and it was rejected because the queries that matter most are the "
        "multi-stage approval and scope-filtered list queries, where the filtering condition is "
        "the security boundary. An object-relational mapper would generate that condition out of "
        "sight or require raw SQL inside it anyway. Parameterised SQL keeps it visible, at a cost "
        "in portability the project does not need.",

        "Server-side sessions were the alternative to tokens, and they revoke instantly, which "
        "tokens do not. They were rejected because the client and the API are separate origins, "
        "where cookie sessions need cross-site configuration and separate protection against "
        "request forgery, whereas tokens do not (Dalimunthe et al., 2023). The revocation "
        "weakness was answered instead: roles are read from the database on every request rather "
        "than trusted from the token, so withdrawal takes effect immediately.",
    ]),

    # ---- 2.4.3 D: Gemini ---------------------------------------------------
    ([
        "Large language models are increasingly used as a natural language interface",
        "The first alternative to the assistant is not another model",
        "Among the hosted models",
    ], [
        "Shi et al. (2025) review the text-to-SQL approach, in which a large language model turns "
        "a question written in ordinary language into a database query, and note that it lowers "
        "the barrier for users who need data they cannot query themselves. That applies directly "
        "here, where staff across several departments need answers about events, requests and "
        "tasks without opening reporting screens belonging to other roles.",

        "The first alternative was no model at all, answering the same questions with more "
        "reporting screens. It was rejected because the number of screens grows with the number "
        "of roles, while the questions staff actually ask cut across them. A self-hosted "
        "open-weight model is stronger on confidentiality and on cost, but it needs a machine "
        "with a graphics processor the project does not have, and is less accurate on this task "
        "(Shi et al., 2025).",

        "Gemini, the OpenAI models and Claude are close enough in capability that the choice "
        "between them is not a capability judgement. Gemini was taken for its documented "
        "development kit at an affordable tier and its support for two credentials with failover "
        "during a live demonstration. The design does not depend on it: the model only proposes a "
        "query, which the guard in Section 4.6.11 validates and constrains to the user's own "
        "access rights before execution.",
    ]),

    # ---- 2.4.4 DBMS --------------------------------------------------------
    ([
        "PostgreSQL is well-suited for managing structured relationships",
        "This choice was made in two parts",
        "Firebase, with its Firestore database",
        "The second part of the choice was where the database runs",
    ], [
        "PostgreSQL suits the multi-departmental workflow because its relational model represents "
        "events, departments, users and tasks directly, its keys and constraints keep the links "
        "between connected records valid, and its transactional behaviour preserves consistency "
        "when several users act on one request at the same time. Comparative studies report "
        "strong query performance against other widely used systems (Salunke & Ouda, 2024; "
        "Urnikienė et al., 2026).",

        "Four alternatives were rejected. MongoDB suits self-contained records, but one event "
        "request refers to a user, a unit, a role assignment, departmental tasks, a venue and a "
        "menu, and a document store would push that integrity into the application. MySQL would "
        "have served, but Salunke and Ouda (2024) report PostgreSQL ahead of it. SQLite is "
        "single-writer and cannot support concurrent approvers. Firestore expresses access as "
        "per-document rules and offers no multi-table transaction, which the audit guarantee "
        "requires.",

        "Where the database runs was decided separately. Hosting it on the development machine is "
        "free but ties institutional data to one computer with no backup, and university "
        "infrastructure would be correct in production but needs a provisioning decision outside "
        "a student project. Supabase supplies the instance, its backups and the connection "
        "pooling, which a single developer is least able to guarantee, while the schema stays "
        "under the project's own versioned migration runner.",
    ]),

    # ---- 2.4.5 operating system -------------------------------------------
    ([
        "Development could equally have been carried out on Linux or macOS",
    ], [
        "Linux or macOS would have served equally for development, and Docker was considered so "
        "that the development and deployment environments would match. Containers were rejected "
        "as disproportionate for one developer with one deployment target, since their isolation "
        "solves a problem that appears when several services and several developers must agree on "
        "an environment. The risk is contained instead, because the Flask application uses no "
        "platform-specific behaviour and every environment reaches the same managed database.",
    ]),

    # ---- 2.4.6 deployment --------------------------------------------------
    ([
        "The platform is deployed as a locally hosted application",
        "Four hosting arrangements were compared",
        "The arrangement adopted splits the decision",
    ], [
        "The Angular client and the Flask API run on one host, the client compiled to static "
        "files and the API served through Gunicorn, while PostgreSQL is hosted by Supabase. The "
        "coordination workload is institutional rather than public, so demand on the application "
        "tier is predictable and needs no elastic scaling, whereas the database carries the data "
        "that must survive, be backed up, and be reachable from every environment.",

        "Four arrangements were rejected. A managed platform such as Render or Railway removes "
        "server administration, but the free tiers suspend an idle instance, and a suspended "
        "instance cannot run the scheduled sweep for reminders and escalation. A cloud virtual "
        "machine, the route Arriesgado et al. (2023) took for a comparable campus system, "
        "transfers patching, certificates and backups onto one student. Orchestrated containers "
        "scale services this system has one of. University hosting needs an institutional "
        "decision.",

        "The arrangement adopted splits the decision along the line between what can be rebuilt "
        "and what cannot. The application tier reconstructs from source in minutes, so it stays "
        "local at no cost, while the database holds the only state that cannot be reconstructed "
        "and is therefore the part placed under managed hosting. Two limits are stated rather "
        "than hidden: uploaded images are written to the host file system, and the scheduled jobs "
        "need a continuously running host. Both are resolved by hosting the application tier, "
        "identified as future work.",
    ]),

    # ---- 2.4.7 browsers ----------------------------------------------------
    ([
        "Cross-browser inconsistencies in how web applications",
        "The alternatives here concern how wide the supported set",
        "Chrome and Firefox were chosen as the pair",
    ], [
        "Choudhary, Prasad and Orso (2012) document how differences between browser "
        "implementations cause the same web application to behave differently, which matters most "
        "for framework-driven interfaces where rendering and JavaScript execution must stay "
        "consistent. Chrome and Firefox are both standards-compliant and available on desktop and "
        "mobile, so staff reach the platform from any device without installing anything.",

        "Three alternatives were rejected. Testing one browser halves the work but, on the "
        "evidence above, produces no evidence about any other engine. Supporting every current "
        "browser including Safari adds a full test pass for each engine, and Safari cannot be "
        "exercised from the Windows machine used here, so its support could be claimed but not "
        "shown. A native application would need a build for each platform and an installation "
        "staff cannot perform.",

        "Chrome and Firefox were chosen as a pair because they run independent engines, Blink and "
        "Gecko. That is what makes the pair meaningful: Chrome together with Microsoft Edge would "
        "cover more users while testing one engine twice. Two independent engines is the smallest "
        "set that gives real evidence of portability, and the largest this project can re-test on "
        "every release without displacing the development it supports.",
    ]),

    # ---- 3.4 pure vs hybrid ------------------------------------------------
    ([
        "The case for a hybrid should not be made without stating its risk",
    ], [
        "A hybrid carries its own risk, since a process that borrows from two models can inherit "
        "the weaknesses of both, planning less than Waterfall requires while iterating less than "
        "Agile expects. Hron and Obwegeser (2018) treat such adaptation as ordinary practice, but "
        "it is only defensible when what is taken from each model is stated in advance. That "
        "condition is applied in the next section, where each part of the process is governed by "
        "one model.",
    ]),

    # ---- 3.5 rejection of the pure methodologies ---------------------------
    ([
        "The three pure alternatives were rejected on specific grounds",
        "A pure Scrum implementation was rejected for a reason of scale",
        "Rapid Application Development came closest to being adopted",
    ], [
        "The three pure alternatives were each rejected in part rather than wholesale. Waterfall "
        "assumes requirements are settled before design begins, the condition Srivastava (2019) "
        "identifies as favouring it, and that did not hold here: the fields a cafeteria request "
        "needs, and the point at which a rejected proposal returns to its applicant, emerged "
        "during implementation. Under pure Waterfall each would have been a change against a "
        "frozen specification. Its documented architecture was kept.",

        "Pure Scrum was rejected on scale rather than on principle. Its stand-ups synchronise "
        "people who cannot see each other's work, its retrospectives address how a group works "
        "together, and its roles divide what to build from how to build it. Kadenic et al. (2023) "
        "tie its benefit to team maturity and to those components. With one developer none of "
        "those problems exist. Short time-boxed iterations and continuous testing were kept.",

        "RAD came closest to being adopted outright, since prototyping is central here. It was "
        "rejected as a complete methodology because it de-emphasises planning, and the parts of "
        "this system that must not change late are the ones it defers: the identity model pairing "
        "user, role and unit, and the transaction binding an audit record to its state change. "
        "Thesing et al. (2021) treat this choice as project-specific, which is the reasoning "
        "applied: prototyping governs the interface, planning governs the architecture.",
    ]),

    # ---- 3.7.1 data collection --------------------------------------------
    ([
        "Three alternatives were considered before the survey was chosen",
        "Analysis of existing documentation was considered as a fourth option",
    ], [
        "Three alternatives were rejected. Semi-structured interviews give a richer account and "
        "allow follow-up, but the question at this stage was whether the difficulties are "
        "widespread, and too few people can be interviewed in the time available to settle that. "
        "Focus groups are hard to assemble across departments within a semester and remove the "
        "anonymity needed when participants criticise their own institution. Direct observation "
        "is the most faithful, but one event cycle runs for weeks across departments a student "
        "cannot enter.",

        "Analysing existing documentation failed for a reason that is itself a finding: there is "
        "no single set of records, because the process is spread across email, spreadsheets and "
        "messaging, the fragmentation stated in Section 1.2.2.1. The survey was the only method "
        "that could reach a group large enough to read as a pattern, anonymously and in time. Its "
        "weakness, that fixed questions miss what the writer did not think to ask, is why the "
        "open-ended question in Section 3.7.4 was included.",
    ]),

    # ---- 5.2 test plan -----------------------------------------------------
    ([
        "Two other approaches to verification were available",
    ], [
        "Two other approaches were available. Automated end-to-end testing through a driven "
        "browser, with Selenium or Cypress, exercises the rules through the interface itself, but "
        "tests of that kind bind to page structure that was still changing, so the effort would "
        "have gone into repairing them rather than extending coverage. Live use by real staff "
        "would give the strongest evidence of all, and requires an institutional authorisation "
        "this project does not hold.",
    ]),
]

# Long paragraphs that are not Justifications but carry a sentence added earlier;
# shortened here so the addition does not lengthen them.
TRIMS = [
    ("Each of those decisions was argued against the alternatives available for it rather than "
     "on its own merits alone, and where the rejected option is the more common industry choice, "
     "the reason it was set aside is stated in terms of what this system requires rather than in "
     "terms of what the tool lacks.",
     "Each of these decisions was argued against the alternatives available for it rather than on "
     "its own merits alone."),

    ("The three pure alternatives were rejected on grounds specific to this project rather than "
     "on general preference: Waterfall because the departmental requirements were not settled "
     "before design began, Scrum because its coordination mechanisms address problems that do not "
     "arise for a single developer, and RAD because it defers the architectural decisions this "
     "system cannot afford to revise late. What each contributes was kept, and Section 3.5 states "
     "which part of the process each of them governs.",
     "The three pure alternatives were rejected on grounds specific to this project, and Section "
     "3.5 states which part of the process each of them still governs."),
]


def set_text(p, text: str) -> None:
    """Write the whole paragraph into its first text run and clear the rest.

    A tab standing before that first run is the body indent and is kept; a tab
    standing after it belonged to text now replaced, and is removed.
    """
    nodes = [n for r in p._p.findall(qn("w:r")) for n in r]
    seen_text = False
    for node in nodes:
        if node.tag == qn("w:t"):
            if not seen_text:
                node.text = text
                node.set(qn("xml:space"), "preserve")
                seen_text = True
            else:
                node.text = ""
        elif node.tag == qn("w:tab") and seen_text:
            node.getparent().remove(node)


def clone_after(p, text: str):
    """A new paragraph after p, carrying p's own paragraph and run properties."""
    new = copy.deepcopy(p._p)
    for attr in (qn("w14:paraId"), qn("w14:textId")):
        new.attrib.pop(attr, None)
    ppr = new.find(qn("w:pPr"))
    for child in list(new):
        if child.tag != qn("w:pPr"):
            new.remove(child)
    run = new.makeelement(qn("w:r"), {})
    rpr = ppr.find(qn("w:rPr")) if ppr is not None else None
    if rpr is not None:
        run.append(copy.deepcopy(rpr))
    if has_indent_tab(p):
        run.append(new.makeelement(qn("w:tab"), {}))
    t = new.makeelement(qn("w:t"), {})
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    new.append(run)
    p._p.addnext(new)
    return new


def has_indent_tab(p) -> bool:
    for node in (n for r in p._p.findall(qn("w:r")) for n in r):
        if node.tag == qn("w:tab"):
            return True
        if node.tag == qn("w:t") and node.text:
            return False
    return False


def find(doc, needle: str):
    hits = [p for p in doc.paragraphs if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(f"anchor matched {len(hits)} paragraphs: {needle!r}")
    return hits[0]


def main() -> int:
    path = sys.argv[1]
    doc = docx.Document(path)

    rewritten = removed = added = 0
    for anchors, texts in BLOCKS:
        paras = [find(doc, a) for a in anchors]
        for para, text in zip(paras, texts):
            set_text(para, text)
            rewritten += 1
        for para in paras[len(texts):]:
            para._p.getparent().remove(para._p)
            removed += 1
        tail = paras[len(texts) - 1] if len(texts) <= len(paras) else paras[-1]
        for text in texts[len(paras):]:
            from docx.text.paragraph import Paragraph
            tail = Paragraph(clone_after(tail, text), doc)
            added += 1
        words = [len(t.split()) for t in texts]
        lines = ", ".join(f"{w / 14:.1f}" for w in words)
        print(f"  {anchors[0][:44]:46s} {len(texts)}p  lines: {lines}")

    trimmed = 0
    for old, new in TRIMS:
        for p in doc.paragraphs:
            if old in p.text:
                set_text(p, p.text.replace(old, new))
                trimmed += 1
                break

    doc.save(path)
    print(f"\n  paragraphs rewritten : {rewritten}")
    print(f"  paragraphs removed   : {removed}")
    print(f"  paragraphs added     : {added}")
    print(f"  sentences trimmed    : {trimmed}")
    over = [(a[0], len(t.split())) for a, ts in BLOCKS for t in ts
            for _ in [0] if len(t.split()) > 98]
    under = [(a[0], len(t.split())) for a, ts in BLOCKS for t in ts
             for _ in [0] if len(t.split()) < 56]
    print(f"  over 7 lines         : {over or 'none'}")
    print(f"  under 4 lines        : {under or 'none'}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
