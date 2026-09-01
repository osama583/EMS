# -*- coding: utf-8 -*-
"""Groups A and B: bring Chapters 1 to 3 into line with a completed project.

Group A repairs Section 1.7, which still described the four-chapter Interim
Report. Group B removes the plan-era voice from Chapters 1 to 3, where the
system is spoken of as something that will be built rather than something that
was built and tested.

Replacement is done across runs rather than by flattening the paragraph, and
text inside a field result is never touched - overwriting a REF field's cached
result loses the sentence the next time Word updates fields.
"""
from __future__ import annotations

import sys

import docx
from docx.oxml.ns import qn

# (locator, old, new) - locator picks the paragraph, old must occur once in it
EDITS = [
    # ── Group A: Section 1.7 ────────────────────────────────────────────
    ("1.7 Overview of the Report",
     "1.7 Overview of the Report",
     "1.7 Overview of the FYP Documentation"),

    ("Chapter 4 — Conclusion",
     "Chapter 4 — Conclusion",
     "Chapter 4 — Design and Implementation"),

    ("The four chapters above describe the structure of this Investigation Report",
     "The four chapters above describe the structure of this Investigation Report. The complete "
     "Final Year Project documentation continues from this foundation into system design and "
     "implementation, testing and evaluation of the delivered platform, and a final conclusion "
     "reflecting on the outcome against the objectives stated in this report.",
     "Together these six chapters document the project from its initial problem definition "
     "through to a delivered and tested platform. Chapters 1 to 3 establish the problem, the "
     "supporting literature and the methodology under which the work was carried out, while "
     "Chapters 4 to 6 present the system as it was built, the evidence obtained from testing it, "
     "and an evaluation of that outcome against the objectives stated in this chapter."),

    # ── Group B1: Section 1.6, benefits were delivered ──────────────────
    ("The implementation of a centralized event coordination platform",
     "is expected to deliver both operational and organizational improvements",
     "delivers both operational and organizational improvements"),
    ("The implementation of a centralized event coordination platform",
     "This section explains the advantages of the proposed system by presenting the expected "
     "improvements in event coordination practices,",
     "This section explains the advantages of the system by presenting the improvements in event "
     "coordination practices evidenced by the acceptance testing reported in Section 5.3.2,"),

    # ── Group B2: Section 3.8.1, author-voice prediction ────────────────
    ("The survey results confirm a 100% compliance rate",
     "the stakeholders who will interact with the proposed event coordination platform",
     "the stakeholders who interact with the event coordination platform"),
    ("data shows a split between students (80%) and staff members (20%)",
     "While students represent the primary end-users who will be requesting and attending events, "
     "the inclusion of staff members is critical.",
     "While students represent the primary end-users who request and attend events, and staff are "
     "the people who work on running them, the inclusion of both groups is critical."),

    # ── Group B3: Section 3.8.1, forecast now answered by Chapter 5 ─────
    ("User intent to use the new platform is very high",
     "User intent to use the new platform is very high, with 73.3% of participants indicating "
     'they are "Likely" or "Very Likely" to adopt the system. This suggests that '
     "the platform will not face significant user resistance and will likely see high engagement "
     "once implemented.",
     "User intent to use the new platform was very high, with 73.3% of participants indicating "
     'they were "Likely" or "Very Likely" to adopt the system. The survey '
     "therefore indicated little user resistance, and the acceptance testing reported in Section "
     "5.3.2 confirmed that expectation once the platform was in use."),

    # ── Group B4: Section 3.8.2, attribute the prediction ───────────────
    ("Operational Efficiency as a Catalyst",
     "By simplifying forms and centralizing approvals, the proposed platform will not only "
     "organize existing events but likely increase the overall vibrancy of campus life.",
     "By simplifying forms and centralizing approvals, respondents expected the platform to "
     "organize existing events and to increase the overall vibrancy of campus life."),

    # ── Group B5: "proposed" where it means the delivered product ───────
    ("This section defines the boundaries",
     "the boundaries of the proposed system", "the boundaries of the system"),
    ("focuses on the development of a centralized",
     "The proposed system focuses", "The system focuses"),
    ("establish the academic and technical foundation",
     "foundation for the proposed centralized, role-based event coordination platform",
     "foundation for the centralized, role-based event coordination platform"),
    ("summary identifying four key research gaps",
     "the proposed system’s solutions", "the system’s solutions"),
    ("digital process management. For the proposed system",
     "For the proposed system, this means", "For this system, this means"),
    ("RBAC is therefore appropriate",
     "For the proposed event coordination platform", "For this event coordination platform"),
    ("directly contributes to SDG Goal 9 by establishing",
     "The proposed centralized", "The centralized"),
    ("is designed as a web-based platform",
     "The proposed system is designed", "The system is designed"),
    ("presented a literature review supporting",
     "supporting the proposed centralized, role-based event coordination platform",
     "supporting the centralized, role-based event coordination platform"),
    ("evaluated based on their alignment with the technical requirements",
     "characteristics of the proposed centralized event coordination platform",
     "characteristics of the centralized event coordination platform"),
    ("primary users and stakeholders",
     "stakeholders of the proposed event coordination platform",
     "stakeholders of the event coordination platform"),
]


def text_nodes(p):
    """Every w:t outside a field result, in document order."""
    out, depth = [], 0
    for r in p._p.findall(qn("w:r")):
        for node in r:
            if node.tag == qn("w:fldChar"):
                kind = node.get(qn("w:fldCharType"))
                if kind == "begin":
                    depth += 1
                elif kind == "end":
                    depth = max(0, depth - 1)
            elif node.tag == qn("w:t") and depth == 0:
                out.append(node)
    return out


def replace_across_runs(p, old: str, new: str) -> bool:
    nodes = text_nodes(p)
    full = "".join(n.text or "" for n in nodes)
    at = full.find(old)
    if at < 0:
        return False
    end, pos, payload = at + len(old), 0, new
    for n in nodes:
        text = n.text or ""
        a, b = pos, pos + len(text)
        pos = b
        if b <= at or a >= end:
            continue
        lo, hi = max(at, a) - a, min(end, b) - a
        n.text = text[:lo] + payload + text[hi:]
        n.set(qn("xml:space"), "preserve")
        payload = ""
    return True


def main() -> int:
    path = sys.argv[1]
    doc = docx.Document(path)

    done = 0
    for locator, old, new in EDITS:
        hits = [p for p in doc.paragraphs if locator in p.text]
        if len(hits) != 1:
            raise SystemExit(f"locator matched {len(hits)}: {locator[:60]!r}")
        if not replace_across_runs(hits[0], old, new):
            raise SystemExit(f"text not found in paragraph: {old[:70]!r}")
        done += 1
        print(f"  {old[:52]:54s} -> {new[:44]}")

    # the project plan calls 1.7 by its old IR name
    rows = 0
    for table in doc.tables:
        for row in table.rows:
            cell = row.cells[0]
            if cell.text.strip() == "1.7 Overview of IR":
                para = cell.paragraphs[0]
                for t in para._p.iter(qn("w:t")):
                    t.text = "1.7 Overview of FYP Doc."
                    break
                rows += 1
    print(f"\n  paragraph edits : {done}")
    print(f"  plan rows renamed: {rows}")

    doc.save(path)
    return 0


if __name__ == "__main__":
    sys.exit(main())
